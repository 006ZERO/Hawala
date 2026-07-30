from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = "commercial/HAWALA_Compliance_OS_Buyer_Pack.docx"
NAVY = "173A35"
GREEN = "2F8F73"
PALE = "EAF4F0"
GRAY = "65736F"
INK = "17211F"
LINE = "D9E2DE"

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Inches(0.8)
sec.left_margin = sec.right_margin = Inches(0.9)
sec.header_distance = sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Aptos"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
for name, size, before, after in [("Heading 1", 17, 16, 7), ("Heading 2", 13, 11, 5), ("Heading 3", 11, 8, 4)]:
    s = styles[name]
    s.font.name = "Aptos Display"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor.from_string(NAVY if name != "Heading 3" else GREEN)
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, (h, width) in enumerate(zip(headers, widths)):
        t.columns[i].width = Inches(width)
        c = t.rows[0].cells[i]
        c.width = Inches(width)
        shade(c, NAVY)
        set_cell_margins(c)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths)):
            cells[i].width = Inches(width)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            if len(t.rows) % 2 == 1:
                shade(cells[i], "F7F9F8")
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(str(value)).font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)

def callout(label, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    shade(c, PALE)
    set_cell_margins(c, 140, 180, 140, 180)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label.upper() + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def page_break():
    doc.add_page_break()

header = sec.header.paragraphs[0]
header.text = "HAWALA COMPLIANCE OS  |  CONFIDENTIAL BUYER PACK"
header.style = styles["Normal"]
header.runs[0].font.size = Pt(8)
header.runs[0].font.bold = True
header.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("Private evaluation material • July 2026").font.size = Pt(8)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(8)
r = p.add_run("HAWALA")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(GREEN)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Compliance OS")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor.from_string(NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(26)
r = p.add_run("A controlled operating layer for broker-led remittance networks")
r.font.size = Pt(16)
r.font.color.rgb = RGBColor.from_string(GRAY)
callout("Current status", "Private, synthetic-data product demonstration. The platform shows durable workflows and control design; sanctions, regulator, payment, and ledger connections are not live.")
doc.add_paragraph("Buyer pack", style="Heading 2")
table(["Purpose", "Audience", "Decision"], [
    ("Product and pilot evaluation", "Regulated remittance providers, banks, central banks, development partners", "Select a bounded discovery and pilot pathway"),
], [2.05, 2.3, 2.15])
doc.add_paragraph("Prepared for confidential commercial discussions. This document is not legal advice, a regulatory approval, an offer of financial services, or evidence of production certification.")

page_break()
doc.add_heading("1. Executive product brief", level=1)
doc.add_paragraph("HAWALA Compliance OS is a compliance and operations platform designed to help licensed or sponsor-supervised remittance brokers record transfers, screen risk, investigate alerts, reconcile obligations, and prepare regulatory evidence without discarding the speed and relationship model that makes broker networks useful.")
doc.add_heading("The buyer problem", level=2)
for x in [
    "Broker-led transfers are often recorded across paper, messaging apps, and disconnected spreadsheets.",
    "Compliance evidence is assembled after the fact, making review slow and regulator engagement difficult.",
    "Settlement positions, prefunding, disputes, and case decisions lack a consistent accountable record.",
    "Existing banking systems can be too costly or operationally heavy for small remittance agents.",
]: bullet(x)
doc.add_heading("What the product demonstrates today", level=2)
table(["Capability", "Demonstrated state", "Production dependency"], [
    ("Transfer and customer records", "Durable workflow with structured records", "Customer data model, migration, and retention approval"),
    ("AML case management", "Rule provenance, evidence context, rationale, human override", "Validated rules/models and licensed screening data"),
    ("Broker onboarding", "License minimization, ownership review, compliance contact, corridors", "Jurisdiction-specific KYB and document verification"),
    ("Regulatory filing", "Draft, human approval, simulated receipt", "Authority-approved schema, credentials, and secure connector"),
    ("Settlement", "Prefunding, netting, reconciliation, dispute record", "Safeguarded account/payment-rail integration"),
], [1.65, 2.35, 2.5])

page_break()
doc.add_heading("2. Product architecture and data flow", level=1)
doc.add_paragraph("The recommended architecture is integration-first and ledger-optional. A conventional auditable database remains the system of record unless a buyer has a proven multi-party governance requirement for a distributed ledger.")
table(["Layer", "Responsibility", "Boundary"], [
    ("User experience", "Mobile-first broker operations and compliance/regulator dashboards", "Role-appropriate views; no unnecessary identity exposure"),
    ("API and workflow", "Transfer, customer, broker, case, filing, settlement, and audit services", "Server-side authorization and validation"),
    ("Risk orchestration", "Rules, list screening, explainability, provenance, analyst decision", "Decision support only; accountable human approval"),
    ("System of record", "Structured records, configuration, filing drafts, audit events", "Retention, residency, backup, encryption defined per buyer"),
    ("External adapters", "Screening provider, identity/KYB, payment rail, regulator interface", "Contracted and tested separately; none live in this demo"),
    ("Optional proof layer", "Shared event proof for governed multi-party use cases", "Not required for settlement and not a substitute for funds movement"),
], [1.25, 2.8, 2.45])
doc.add_heading("Illustrative transaction flow", level=2)
for i, x in enumerate([
    "Broker records sender, recipient, corridor, purpose, and transfer value.",
    "Policy engine evaluates configured rules and licensed external data when connected.",
    "Low-risk transfers proceed; alerts enter a human-owned compliance case.",
    "Analyst reviews evidence, records rationale, and can apply a reasoned override.",
    "Approved transfer updates broker obligations and the next reconciliation cycle.",
    "Authorized staff prepare a filing draft; transmission occurs only through an approved connector.",
], 1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(x)

page_break()
doc.add_heading("3. Proposed pilot", level=1)
callout("Recommendation", "Begin with one sponsor, 5–10 brokers, one corridor, and synthetic or masked data. Do not begin with live regulator submission or autonomous risk decisions.")
table(["Phase", "Duration", "Outputs", "Exit gate"], [
    ("0. Discovery", "2 weeks", "Legal perimeter, corridor, roles, data map, success baseline", "Sponsor approves scope and assumptions"),
    ("1. Configure", "3–4 weeks", "Sandbox, workflows, policy pack, synthetic migration, training", "Security and UAT readiness"),
    ("2. Shadow pilot", "4–6 weeks", "Parallel records, screening comparison, case and reconciliation evidence", "No funds movement; accuracy accepted"),
    ("3. Controlled live pilot", "6–8 weeks", "Limited broker cohort and value caps; supervised operations", "Regulatory/sponsor approvals obtained"),
    ("4. Scale decision", "2 weeks", "Results, gaps, production plan, commercial proposal", "Joint go/no-go"),
], [1.1, 0.85, 3.0, 1.55])
doc.add_heading("Pilot success measures", level=2)
for x in [
    "At least 95% of required transfer records captured with complete mandatory fields.",
    "100% of case decisions attributed with rationale and evidence provenance.",
    "Reconciliation exceptions identified and assigned within one business day.",
    "Measured screening precision/recall compared with the sponsor-approved baseline.",
    "No external filing, funds movement, or data sharing outside the approved pilot boundary.",
]: bullet(x)

page_break()
doc.add_heading("4. Commercial model and indicative pricing", level=1)
doc.add_paragraph("Pricing is an indicative commercial framework for buyer discussion, excluding taxes, third-party data, payment rails, identity checks, hosting-specific compliance, legal advice, and regulator certification.")
table(["Component", "Indicative price", "Includes"], [
    ("Discovery and pilot design", "USD 15,000–25,000", "Scope, controls, architecture, success plan"),
    ("Configured sandbox pilot", "USD 40,000–75,000", "Configuration, onboarding, training, support, evaluation"),
    ("Production implementation", "USD 120,000–250,000", "Integration, migration, security hardening, UAT, launch"),
    ("Platform subscription", "USD 4,000–12,000/month", "Core platform, support tier, agreed usage band"),
    ("Transaction usage", "USD 0.10–0.60/transfer", "Volume-based platform usage; third-party fees separate"),
], [1.7, 1.45, 3.35])
doc.add_heading("Commercial protections", level=2)
for x in [
    "Milestone billing tied to accepted deliverables and explicit exit gates.",
    "Change control for jurisdiction, corridor, integration, or data-scope expansion.",
    "No claim that platform pricing alone guarantees an end-customer fee below 2%.",
    "Buyer retains accountability for licensing, policy approval, filing, and funds safeguarding.",
]: bullet(x)

page_break()
doc.add_heading("5. Security, privacy, and control summary", level=1)
table(["Control domain", "Demonstrated", "Required before production"], [
    ("Identity and access", "Workspace identity; server-side Administrator, Compliance Officer, Operator, Auditor roles", "Buyer IAM/SSO, provisioning, MFA policy, periodic access review"),
    ("Auditability", "Durable attributed workflow events and decision rationale", "Immutability design, export, monitoring, retention validation"),
    ("Data minimization", "License number limited to last four characters in broker workflow", "Field-level inventory, lawful basis, deletion and subject-rights process"),
    ("Encryption and secrets", "Hosting-provider managed baseline", "Key ownership, rotation, secrets vault, TLS and at-rest evidence"),
    ("Resilience", "Not certified in demo", "Backups, restoration tests, RTO/RPO, incident response and BCP"),
    ("Model governance", "Rule/model versions and evidence provenance shown", "Validation, drift monitoring, bias testing, change approval and rollback"),
], [1.35, 2.4, 2.75])
callout("Security caveat", "The demo is privately accessible, but it is not a production security certification. Data residency, penetration testing, encryption evidence, subprocessor review, and contractual controls remain buyer-specific work.")

page_break()
doc.add_heading("6. Regulatory and operating assumptions", level=1)
for x in [
    "The buyer or sponsoring institution holds all required permissions and determines whether participating brokers may operate.",
    "Country-by-country legal analysis is required; terminology such as broker, agent, hawaladar, payment service provider, and money service business is not interchangeable.",
    "AML screening supports—not replaces—risk-based judgment. A human remains accountable for escalation, override, and filing approval.",
    "Suspicious transaction reports are prepared and submitted only under the buyer's approved process and authority interface.",
    "Tax visibility does not mean automatic taxation. Collection, use, disclosure, and retention require a lawful and documented basis.",
    "Settlement requires safeguarded funds, prefunding or approved credit arrangements, reconciliation, dispute handling, and licensed payment rails.",
    "Sanctions and PEP data must come from licensed, current, contractually permitted sources with documented update timing.",
    "Distributed ledger technology is optional and requires a governance case; it neither moves fiat funds nor creates regulatory permission.",
]: bullet(x)
doc.add_heading("Synthetic-data disclaimer", level=2)
doc.add_paragraph("All persons, transactions, brokers, alerts, corridors, receipts, sanctions/PEP matches, regulator connections, balances, and performance values shown in the demonstration are synthetic or illustrative unless a signed pilot data schedule states otherwise. Demo outcomes must not be used for real customer decisions, reporting, funds movement, or compliance conclusions.")

page_break()
doc.add_heading("7. Buyer decision checklist", level=1)
table(["Decision area", "Buyer input needed"], [
    ("Sponsor and legal perimeter", "Licensed entity, target jurisdiction, pilot authority, accountable executive"),
    ("Pilot corridor", "Origin/destination, currency, projected volume, broker cohort"),
    ("Data", "Permitted fields, masking, residency, retention, migration source"),
    ("Risk", "Policy owner, rules, screening providers, validation baseline, approval thresholds"),
    ("Operations", "Prefunding model, value caps, reconciliation timing, disputes, safeguarding"),
    ("Technology", "IAM, integrations, hosting controls, monitoring, environments"),
    ("Regulatory reporting", "Schema, approval chain, secure channel, test environment"),
    ("Commercial", "Pilot budget, procurement, contracting, third-party costs, scale criteria"),
], [1.75, 4.75])
callout("Next step", "A 90-minute buyer workshop should confirm the legal perimeter, one corridor, one sponsor, pilot data boundaries, and measurable exit gates. A tailored statement of work can then be produced without overstating readiness.")
doc.add_paragraph("Document control: Buyer Pack v1.0 • Prepared July 2026 • Confidential evaluation material")

doc.core_properties.title = "HAWALA Compliance OS — Confidential Buyer Pack"
doc.core_properties.subject = "Commercial product, pilot, pricing, architecture, security, and regulatory assumptions"
doc.core_properties.author = "HAWALA Compliance OS"
doc.save(OUT)
print(OUT)
